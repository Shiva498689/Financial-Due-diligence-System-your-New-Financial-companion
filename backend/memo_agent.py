from __future__ import annotations
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
import asyncio
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from google import genai
from google.genai.types import GenerateContentConfig
try:
    from chart_generator import generate_all_charts, generate_all_charts_async
    CHARTS_ENABLED = True
except ImportError:
    CHARTS_ENABLED = False
    def generate_all_charts(state, out_dir):
        return {}
    async def generate_all_charts_async(state, out_dir):
        return {}
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
client = genai.Client(api_key=GEMINI_API_KEY) if GEMINI_API_KEY else None
OUTPUTS_DIR = Path(__file__).parent / "outputs"
LOGO_PATH = Path(__file__).parent / "assets" / "wallstreet_dd_logo.png"
GOLD_RGB    = RGBColor(0xC9, 0xA8, 0x4C)  
WHITE_RGB   = RGBColor(0xFF, 0xFF, 0xFF)   
DARK_RGB    = RGBColor(0x1A, 0x1A, 0x1A)   
NAVY_RGB    = RGBColor(0x1B, 0x3A, 0x6B)   
LGREY_RGB   = RGBColor(0xF4, 0xF6, 0xF9)  
XL_NAVY     = "1B3A6B"
XL_GOLD     = "C9A84C"
XL_GOLD_BG  = "C9A84C"
XL_WHITE    = "FFFFFF"
XL_LGREY    = "F4F6F9"
XL_DARK     = "1A1A1A"
XL_GREEN    = "1E8B4C"
XL_AMBER    = "E67E22"
XL_RED      = "C0392B"
XL_SECTION  = "C9A84C"   
XL_SEC_TXT  = "1A1A1A"   
XL_LIGHT_GOLD = "FFF8E7"  
METRIC_LABELS: Dict[str, tuple] = {
    "Revenue":                    ("Total Revenue",              "Total money the company earned from selling products or services"),
    "CostOfRevenue":              ("Cost of Revenue (COGS)",     "Direct costs to produce what was sold (materials, manufacturing)"),
    "GrossProfit":                ("Gross Profit",               "Revenue minus the direct cost of production — money left before operating expenses"),
    "ResearchAndDevelopment":     ("R&D Spending",               "Money spent on developing new products or improving existing ones"),
    "SellingGeneralAndAdministrative": ("SG&A Expenses",         "Overhead costs: salaries, marketing, admin — not directly tied to production"),
    "DepreciationAndAmortization":("Depreciation & Amortization","Non-cash accounting charge for wear and tear on assets over time"),
    "OperatingExpenses":          ("Total Operating Expenses",   "All costs to run the business day-to-day before interest and taxes"),
    "OperatingIncome":            ("Operating Income (EBIT)",    "Profit from core business operations — before interest costs and taxes"),
    "InterestExpense":            ("Interest Paid on Debt",      "Cost of borrowing money — paid to lenders on outstanding loans or bonds"),
    "IncomeTaxExpense":           ("Income Tax Paid",            "Taxes paid to governments on the company's profits"),
    "NetIncome":                  ("Net Profit (Bottom Line)",   "Final profit after ALL costs, interest, and taxes are deducted from revenue"),
    "CashAndEquivalents":         ("Cash & Short-term Savings",  "Money the company has immediately available — cash and near-cash assets"),
    "ShortTermInvestments":       ("Short-term Investments",     "Liquid investments that can be converted to cash within 12 months"),
    "AccountsReceivable":         ("Money Owed to Company",      "Revenue already earned but not yet collected from customers"),
    "Inventory":                  ("Unsold Inventory",           "Goods produced but not yet sold — sits on shelves or in warehouses"),
    "CurrentAssets":              ("Short-term Assets",          "Assets expected to be used or converted to cash within 12 months"),
    "PropertyPlantAndEquipment":  ("Property, Plant & Equipment","Physical assets: buildings, machines, vehicles, and infrastructure"),
    "TotalAssets":                ("Total Assets",               "Everything the company owns — both short-term and long-term assets combined"),
    "AccountsPayable":            ("Money Owed to Suppliers",    "Bills the company has received but not yet paid to its vendors"),
    "ShortTermDebt":              ("Short-term Borrowings",      "Loans and debt that must be repaid within the next 12 months"),
    "CurrentLiabilities":         ("Short-term Obligations",     "All financial obligations due within the next 12 months"),
    "LongTermDebt":               ("Long-term Debt",             "Borrowings due in more than 12 months — bonds, bank loans, etc."),
    "TotalLiabilities":           ("Total Debt & Obligations",   "Everything the company owes — both short-term and long-term combined"),
    "RetainedEarnings":           ("Accumulated Retained Profits","Total profits kept inside the company over its entire history (not paid as dividends)"),
    "ShareholdersEquity":         ("Shareholders Equity (Book Value)", "What shareholders own: Total Assets minus Total Liabilities"),
    "OperatingCashFlow":          ("Cash from Operations",       "Actual cash generated from running the core business — not accounting profit"),
    "CapitalExpenditures":        ("Capital Expenditure (CapEx)","Cash spent buying or upgrading physical assets like factories and equipment"),
    "FreeCashFlow":               ("Free Cash Flow (FCF)",       "Cash left over after CapEx — money available to pay dividends, buy back shares, or invest"),
    "GrossMargin":                ("Gross Profit Margin (%)",    "What percentage of each revenue dollar is left after direct production costs"),
    "OperatingMargin":            ("Operating Profit Margin (%)", "What percentage of each revenue dollar remains after all operating costs"),
    "NetProfitMargin":            ("Net Profit Margin (%)",      "What percentage of each revenue dollar becomes final profit"),
    "CurrentRatio":               ("Current Ratio",              "Can the company pay its short-term bills? Above 1.0x means yes. Above 2.0x is strong"),
    "DebtToEquity":               ("Debt-to-Equity Ratio",       "How much debt vs. owner equity is used to finance the business. Higher = more financial risk"),
    "ReturnOnAssets":             ("Return on Assets (ROA %)",   "How efficiently the company uses its assets to generate profit. Higher is better"),
    "ReturnOnEquity":             ("Return on Equity (ROE %)",   "Profit generated for every dollar of shareholder investment. Higher is better"),
}
FRAUD_MODEL_LABELS: Dict[str, Dict] = {
    "piotroski": {
        "name": "Piotroski F-Score",
        "what_it_is": "A 9-point checklist that scores a company's financial strength across profitability, debt, and operating efficiency.",
        "how_to_read": "Score 7–9: Financially strong. Score 4–6: Mixed signals. Score 0–3: Financially weak.",
        "field": "piotroski_f_score",
    },
    "beneish": {
        "name": "Beneish M-Score",
        "what_it_is": "A mathematical model that detects whether a company may be manipulating (inflating) its reported earnings.",
        "how_to_read": "Below -2.5: Unlikely manipulator. Between -2.5 and -1.5: Grey zone. Above -1.5: Possible earnings manipulation.",
        "field": "beneish_m_score",
    },
    "ohlson": {
        "name": "Ohlson O-Score (Bankruptcy Probability)",
        "what_it_is": "A model that calculates the probability a company will go bankrupt within two years based on financial ratios.",
        "how_to_read": "Below 30%: Low distress risk. 30–60%: Moderate concern. Above 60%: High bankruptcy risk.",
        "field": "ohlson_o_score_probability",
    },
    "merton": {
        "name": "Merton Distance to Default",
        "what_it_is": "Measures how far the company's asset value is from the point where it cannot pay its debts — in units of standard deviation.",
        "how_to_read": "Above 2.5: Safe. 1.0–2.5: Moderate credit risk. Below 1.0: Elevated default risk.",
        "field": "merton_distance_to_default",
    },
}
REPORT_SYSTEM_PROMPT = """You are the presentation engine of an institutional financial due-diligence platform.
Be professional. Do not use emojis anywhere. Write clearly for a reader who is not an accountant.
IMPORTANT RULES:
- Do NOT perform any new financial analysis.
- Do NOT invent numbers or estimates.
- Use ONLY the supplied JSON data.
- Explain every financial term in plain, simple English the first time you use it.
- Never print raw numbers like 402836000000 — always format as $402.8B, $1.69T, 59.6%, etc.
Output a single valid JSON object with exactly these keys:
{
  "executive_summary": "3-4 sentence plain-English summary of the company's financial health",
  "overall_risk_rating": "HEALTHY / MODERATE RISK / HIGH RISK — one line",
  "investment_recommendation": "Clear buy/hold/sell recommendation with one-paragraph rationale",
  "business_overview": "What this company does and how it makes money — written for a non-accountant",
  "quantitative_narrative": "2-3 paragraph plain-English explanation of what the KPI scores collectively mean",
  "risk_summary": {
    "severe_risks": ["Plain-English description of each severe risk"],
    "moderate_risks": ["Plain-English description of each moderate risk"],
    "low_risks": ["Plain-English description of each low risk or positive factor"]
  },
  "valuation_narrative": "Plain-English explanation of what the DCF and Monte Carlo values mean for an investor",
  "missing_data": ["List ONLY narrative questions that explicitly failed as unavailable. Do NOT list missing financial metrics (e.g. R&D, COGS, Revenue, etc.) as they are handled elsewhere."],
  "risk_flags": ["Each flag as a plain-English sentence describing the concern"]
}
Return ONLY valid JSON. No markdown fences. No raw JSON blobs in any value field.
All numbers must be formatted as human-readable strings: $1.23B, 5.4%, 6/9, etc."""

async def generate_memo_async(state: Dict[str, Any]) -> Dict[str, str]:
    ticker = state.get("ticker", "UNKNOWN").upper()
    out_dir = OUTPUTS_DIR / ticker
    out_dir.mkdir(parents=True, exist_ok=True)
    charts_dir = out_dir / "charts"
    async def _empty_charts():
        return {}

    charts_task = (
        generate_all_charts_async(state, charts_dir)
        if CHARTS_ENABLED
        else _empty_charts()
    )
    structured_task = _call_gemini_for_structured_report_async(state)
    charts, structured = await asyncio.gather(charts_task, structured_task)
    markdown_report = await _call_gemini_for_markdown_async(state, charts)
    docx_path  = out_dir / f"{ticker}_due_diligence.docx"
    excel_path = out_dir / f"{ticker}_dcf_model.xlsx"

    await asyncio.gather(
        asyncio.to_thread(_build_docx,  structured, state, ticker, docx_path, charts),
        asyncio.to_thread(_build_excel, state, ticker, excel_path),
    )

    return {
        "docx_path":       str(docx_path),
        "excel_path":      str(excel_path),
        "markdown_report": markdown_report,
    }
def _build_report_payload(state: Dict[str, Any]) -> dict:
    financial_keys = [
        "revenue_latest", "revenue_previous", "cogs_latest", "cogs_previous",
        "gross_profit_latest", "gross_profit_previous", "sga_expenses_latest", "sga_expenses_previous",
        "depreciation_amortization_latest", "depreciation_amortization_previous",
        "net_income_latest", "net_income_previous", "net_income_continuing_ops_latest", "net_income_continuing_ops_previous",
        "operating_cash_flow_latest", "operating_cash_flow_previous", "capex_latest", "capex_previous",
        "current_assets_latest", "current_assets_previous", "current_liabilities_latest", "current_liabilities_previous",
        "cash_and_equivalents_latest", "cash_and_equivalents_previous", "receivables_latest", "receivables_previous",
        "gross_ppe_latest", "gross_ppe_previous", "total_assets_latest", "total_assets_previous",
        "total_liabilities_latest", "total_liabilities_previous", "long_term_debt_latest", "long_term_debt_previous",
        "short_term_debt_latest", "short_term_debt_previous"
    ]
    return {
        "ticker": state.get("ticker"),
        "risk_report": state.get("risk_report"),#correct
        "quant_analysis": state.get("quant_analysis"),#correct
        "narrative_analysis": state.get("narrative_analysis"),#narrative analysis
        "company_description": state.get("company_description"),#done
        "financial_metrics": {k: state.get(k) for k in financial_keys},
        "market_metrics": {
            "current_equity_price": state.get("current_equity_price"),
            "market_capitalization": state.get("market_capitalization"),
            "historical_equity_volatility_252d": state.get("historical_equity_volatility_252d"),
            "risk_free_rate": state.get("risk_free_rate"),
            "gnp_deflator": state.get("gnp_deflator"),
        },
        "risk_models": {
            "piotroski_f_score": state.get("piotroski_f_score"),
            "beneish_m_score": state.get("beneish_m_score"),
            "ohlson_o_score_probability": state.get("ohlson_o_score_probability"),
            "merton_distance_to_default": state.get("merton_distance_to_default"),
        },
        "valuation": {
            "deterministic_dcf_value": state.get("deterministic_dcf_value"),
            "monte_carlo_p10_floor": state.get("monte_carlo_p10_floor"),
            "monte_carlo_p50_median": state.get("monte_carlo_p50_median"),
            "monte_carlo_p90_ceiling": state.get("monte_carlo_p90_ceiling"),
        },
    }



async def _call_gemini_for_structured_report_async(state: Dict[str, Any]) -> dict:
    payload = _build_report_payload(state)
    response = await client.aio.models.generate_content(
        model="gemini-3.1-flash-lite",
        contents=json.dumps(payload, indent=2),
        config=GenerateContentConfig(
            system_instruction=REPORT_SYSTEM_PROMPT,
            temperature=0.1,
        ),
    )
    raw = response.text.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    try:        return json.loads(raw)
    except json.JSONDecodeError:
        return {"executive_summary": raw, "missing_data": [], "risk_flags": []}
def _call_gemini_for_markdown(state: Dict[str, Any], charts: dict = None) -> str:
    if charts is None:
        charts = {}
    ticker  = state.get("ticker", "UNKNOWN").upper()
    risk    = state.get("risk_report", {})
    overall = risk.get("overall_assessment", {}).get("final_verdict", "N/A")
    avg_risk_score = risk.get("overall_assessment", {}).get("average_risk_score")
    piotroski = state.get("piotroski_f_score", "N/A")
    beneish   = state.get("beneish_m_score", "N/A")
    ohlson    = state.get("ohlson_o_score_probability")
    merton    = state.get("merton_distance_to_default", "N/A")
    dcf   = state.get("deterministic_dcf_value")
    p10   = state.get("monte_carlo_p10_floor")
    p50   = state.get("monte_carlo_p50_median")
    p90   = state.get("monte_carlo_p90_ceiling")
    mktcap = state.get("market_capitalization")
    price  = state.get("current_equity_price")
    rev    = state.get("revenue_latest")
    ni     = state.get("net_income_latest")
    ocf    = state.get("operating_cash_flow_latest")
    rev_p  = state.get("revenue_previous")
    ni_p   = state.get("net_income_previous")
    gp     = state.get("gross_profit_latest")
    oi     = state.get("operating_income_latest")
    capex  = state.get("capex_latest")
    fcf_v  = (ocf or 0) - abs(capex or 0)
    equity = state.get("shareholders_equity_latest")
    debt   = state.get("long_term_debt_latest")
    def fmt(v):
        if v is None: return "N/A"
        v = float(v)
        if abs(v) >= 1e12: return f"${v/1e12:.2f}T"
        if abs(v) >= 1e9:  return f"${v/1e9:.1f}B"
        if abs(v) >= 1e6:  return f"${v/1e6:.1f}M"
        return f"${v:,.2f}"
    def fmt_pct(v):
        if v is None: return "N/A"
        f = float(v)
        return f"{f*100:.4f}%" if abs(f) < 0.01 else f"{f*100:.1f}%"
    def yoy(latest, prior):
        if latest and prior and float(prior) != 0:
            return f"{(float(latest)-float(prior))/abs(float(prior))*100:+.1f}%"
        return "N/A"
    def margin(num, den):
        if num and den and float(den) != 0:
            return f"{float(num)/float(den)*100:.1f}%"
        return "N/A"
    summary_prompt = f"""Write a structured 5-paragraph executive summary for {ticker} based on the data below.
Each paragraph should cover exactly one topic: (1) company financial performance, (2) profitability and margins,
(3) cash flow and balance sheet health, (4) risk model findings, (5) valuation vs market price.
Data:
- Overall risk verdict: {overall} (composite risk score: {avg_risk_score})
- Revenue: {fmt(rev)} (year-on-year change: {yoy(rev, rev_p)})
- Net Profit: {fmt(ni)} (year-on-year change: {yoy(ni, ni_p)})
- Gross Profit: {fmt(gp)}, Gross Margin: {margin(gp, rev)}
- Operating Income: {fmt(oi)}, Operating Margin: {margin(oi, rev)}
- Net Profit Margin: {margin(ni, rev)}
- Operating Cash Flow (actual cash from running the business): {fmt(ocf)}
- Capital Expenditure (money spent on factories/equipment): {fmt(capex)}
- Free Cash Flow (cash remaining after CapEx): {fmt(fcf_v)}
- Shareholders Equity (what owners actually own): {fmt(equity)}
- Long-term Debt: {fmt(debt)}
- Piotroski F-Score (financial strength, scored 0-9, higher is better): {piotroski}/9
- Beneish M-Score (earnings quality detector, below -2.5 means unlikely to be manipulating): {beneish}
- Ohlson O-Score (probability of bankruptcy within 2 years): {fmt_pct(ohlson)}
- Merton Distance to Default (how many standard deviations away from being unable to pay debts): {merton}
- DCF Intrinsic Value (our discounted cash flow model estimate of fair value): {fmt(dcf)}
- Monte Carlo Valuation Range — Bear Case: {fmt(p10)}, Base Case: {fmt(p50)}, Bull Case: {fmt(p90)}
- Current Market Capitalization: {fmt(mktcap)}
- Stock Price: ${price}
Rules:
- No emojis, no bullet points, no markdown headers inside paragraphs
- Write for someone who is NOT an accountant — explain every financial term simply in parentheses when first used
- Do NOT just repeat the numbers — synthesise them into meaningful insight sentences
- Be specific about what these numbers mean for an investor
- Each paragraph should be 3-4 sentences
- Total length: 300-380 words
- End with a clear one-sentence investment stance (Buy / Hold / Sell with brief reason)"""
    response = client.models.generate_content(
        model="gemini-3.1-flash-lite",
        contents=summary_prompt,
        config=GenerateContentConfig(temperature=0.12, max_output_tokens=900),
    )
    prose = response.text.strip()
    def _chart_img(chart_name: str, alt: str, title: str = "") -> str:
        """Return markdown img tag using backend URL. Empty string if chart missing."""
        chart_path = charts.get(chart_name)
        if not chart_path or not Path(str(chart_path)).exists():
            return ""
        fname = Path(str(chart_path)).name
        url = f"/outputs/{ticker}/charts/{fname}"
        title_attr = f' "{title}"' if title else ""
        return f"![{alt}]({url}{title_attr})\n"
    chart_kpi      = _chart_img("kpi_health_scores",   "KPI Health Scores — AI Audit",    "KPI Health Scores (10 = Excellent)")
    chart_revenue  = _chart_img("revenue_breakdown",    "Income Statement Comparison",     "Latest vs Previous Year")
    chart_margins  = _chart_img("margin_trends",        "Profitability Margins",           "Latest vs Previous Year")
    chart_dcf      = _chart_img("dcf_waterfall",        "DCF Valuation Range",             "Monte Carlo Scenarios vs Market Cap")
    chart_radar    = _chart_img("risk_radar",           "Risk Model Radar",                "All 5 Models — 10 = Safest")
    chart_donut    = _chart_img("balance_sheet_pie",    "Balance Sheet Composition",       "Asset Breakdown")
    chart_risk_fc  = _chart_img("risk_flowchart",      "Risk Assessment Flowchart",        "5 Models → Overall Verdict")
    chart_val_fc   = _chart_img("valuation_flowchart", "Valuation Decision Flowchart",     "DCF Scenarios → Market Cap Comparison")
    _verdict      = overall
    _verdict_stl  = "fill:#27AE60,color:#fff" if "HEALTHY" in _verdict else (
                    "fill:#E74C3C,color:#fff" if "HIGH" in _verdict else "fill:#E67E22,color:#fff")
    p_flag   = risk.get("piotroski", {}).get("risk", "N/A")
    m_flag   = risk.get("beneish",   {}).get("risk", "N/A")
    o_flag   = risk.get("ohlson",    {}).get("risk", "N/A")
    d_flag   = risk.get("merton",    {}).get("risk", "N/A")
    dcf_flag = risk.get("dcf_risk",  {}).get("risk", "N/A")
    def _rns(flag):
        if "LOW"  in flag: return "fill:#1E8B4C,color:#fff"
        if "HIGH" in flag: return "fill:#C0392B,color:#fff"
        return "fill:#E67E22,color:#fff"
    mermaid_risk = f"""```mermaid
flowchart TD
    A["{ticker} — Risk Assessment"] --> B["Piotroski<br/>{piotroski}/9<br/>{p_flag}"]
    A --> C["Beneish<br/>{beneish}<br/>{m_flag}"]
    A --> D["Ohlson<br/>{fmt_pct(ohlson)}<br/>{o_flag}"]
    A --> E["Merton<br/>{merton}<br/>{d_flag}"]
    A --> F["DCF Risk<br/>{dcf_flag}"]
    B --> VERDICT["{_verdict}"]
    C --> VERDICT
    D --> VERDICT
    E --> VERDICT
    F --> VERDICT
    style A fill:#1B3A6B,color:#F0D080,stroke:#C9A84C
    style B {_rns(p_flag)}
    style C {_rns(m_flag)}
    style D {_rns(o_flag)}
    style E {_rns(d_flag)}
    style F {_rns(dcf_flag)}
    style VERDICT {_verdict_stl}
```"""
    _over    = mktcap and dcf and float(mktcap) > float(dcf) * 1.1
    _under   = mktcap and dcf and float(mktcap) < float(dcf) * 0.9
    _vlabel  = "Potentially Overvalued" if _over else ("Potentially Undervalued" if _under else "Fairly Valued")
    _vstl    = "fill:#E74C3C,color:#fff" if _over else ("fill:#27AE60,color:#fff" if _under else "fill:#E67E22,color:#fff")
    mermaid_valuation = f"""```mermaid
flowchart LR
    P10["Bear Case<br/>{fmt(p10)}"]
    P50["Base Case<br/>{fmt(p50)}"]
    P90["Bull Case<br/>{fmt(p90)}"]
    DCF[["DCF Estimate<br/>{fmt(dcf)}"]]
    MC[("Market Cap<br/>{fmt(mktcap)}")]
    VRD["{_vlabel}"]
    P10 --> DCF
    P50 --> DCF
    P90 --> DCF
    DCF --> VRD
    MC --> VRD
    style P10 fill:#C0392B,color:#fff
    style P50 fill:#C9A84C,color:#000
    style P90 fill:#27AE60,color:#fff
    style DCF fill:#1B3A6B,color:#F0D080,stroke:#C9A84C
    style MC  fill:#142233,color:#E8EDF2,stroke:#8E9BAD
    style VRD {_vstl}
```"""
    _pi_txt = "Strong" if int(piotroski or 0) >= 7 else ("Moderate" if int(piotroski or 0) >= 4 else "Weak")
    _be_txt = "Low manipulation risk" if float(beneish or 0) < -2.5 else "Requires closer review"
    _me_txt = f"{float(merton or 0):.2f} std dev" if merton and merton != "N/A" else "N/A"
    metrics_table = f"""
| Metric | Value | What it means |
|--------|-------|---------------|
| Revenue (total sales) | {fmt(rev)} | Money earned from all business activities |
| Gross Profit | {fmt(gp)} | Revenue minus direct production costs |
| Net Profit (bottom line) | {fmt(ni)} | Money left after paying ALL costs and taxes |
| Free Cash Flow | {fmt(fcf_v)} | Cash available after maintaining assets |
| Operating Cash Flow | {fmt(ocf)} | Actual cash generated from running the business |
| Stock Price | ${price} | Current market price per share |
| Market Capitalization | {fmt(mktcap)} | Total market value of all shares combined |
| DCF Intrinsic Value | {fmt(dcf)} | Model estimate of what the company is worth |
| Piotroski F-Score | {piotroski}/9 | Financial health score — {_pi_txt} |
| Beneish M-Score | {beneish} | Earnings quality — {_be_txt} |
| Ohlson Bankruptcy Risk | {fmt_pct(ohlson)} | Probability of financial distress in 2 years |
| Merton Default Distance | {_me_txt} | Distance from insolvency (higher = safer) |
"""
    _verdict_badge = ("> **OVERALL RISK VERDICT**\n"
                      f"> **{_verdict}**\n"
                      f"> Composite risk score: {avg_risk_score:.2f}/3.0 "
                      f"(1.0 = fully safe, 3.0 = high risk)\n"
                      if avg_risk_score else
                      f"> **OVERALL RISK VERDICT: {_verdict}**\n")

    return f"""# {ticker} — Due Diligence Summary

{_verdict_badge}

{prose}

---

## Key Metrics at a Glance
{metrics_table}

---

## Revenue & Profitability

{chart_revenue}{chart_margins}

---

## KPI Health Scores

{chart_kpi}

---

## Valuation Analysis

{chart_dcf}

**Valuation Decision — How the numbers compare:**

{mermaid_valuation}

{chart_val_fc}

---

## Risk Assessment

{chart_radar}

**Risk Model Breakdown — All 5 Models:**

{mermaid_risk}

{chart_risk_fc}

---

## Balance Sheet Composition

{chart_donut}

---

> **Full Report Available**
>
> The downloadable Word document contains the complete analysis:
> - All 30 SEC filing review questions answered in detail
> - Full income statement, balance sheet, and cash flow tables
> - All 6 charts with detailed written commentary
> - Fraud detection model deep-dive (Piotroski, Beneish, Ohlson, Merton)
> - Full DCF valuation model with assumptions
> - The accompanying Excel workbook contains the complete DCF model with live formulas
"""
async def _call_gemini_for_markdown_async(state: Dict[str, Any], charts: dict = None) -> str:
    import asyncio
    if charts is None:
        charts = {}
    ticker = state.get("ticker", "UNKNOWN").upper()
    result = await asyncio.to_thread(_call_gemini_for_markdown, state, charts)
    return result
def _build_docx(structured: dict, state: Dict[str, Any], ticker: str, out_path: Path, charts: dict = None) -> None:
    if charts is None:
        charts = {}
    doc = Document()
    _configure_document(doc)
    _add_title_page(doc, ticker, structured)
    doc.add_page_break()
    _heading(doc, "1.  Executive Summary")
    _body(doc, structured.get("executive_summary", "Not available."))
    _heading(doc, "2.  Overall Risk Rating")
    _risk_banner(doc, structured.get("overall_risk_rating", "N/A"))
    _heading(doc, "3.  Investment Recommendation")
    _body(doc, structured.get("investment_recommendation", "Not available."))
    _heading(doc, "4.  Business Overview")
    _body(doc, structured.get("business_overview", "Not available."))
    _heading(doc, "5.  Financial Statements — Latest vs. Previous Year")
    _note(doc, "All figures sourced directly from SEC EDGAR filings. Values shown in millions or billions.")
    _financial_statements_table(doc, state)

   
    _embed_chart(doc, charts.get("revenue_breakdown"), "Income Statement: Latest vs Previous Year ($B)")

    # Chart: Margin Trends
    _embed_chart(doc, charts.get("margin_trends"), "Profitability Margins: Latest vs Previous Year")

    # Section 6 — Quantitative Analysis (all scored KPIs from quant engine)
    _heading(doc, "6.  Quantitative Analysis — Scored Financial Metrics")
    _note(doc, "Each metric below was extracted from SEC EDGAR, scored by our AI engine on a scale of 1–10, and flagged if a risk is detected.")
    _body(doc, structured.get("quantitative_narrative", ""))
    _kpi_table(doc, state)

    # Chart: KPI Health Scores
    _embed_chart(doc, charts.get("kpi_health_scores"), "KPI Health Score Summary — AI Audit Results (10 = Excellent)")

    # Section 7 — Risk Assessment
    _heading(doc, "7.  Risk Assessment")
    _risk_assessment_section(doc, structured.get("risk_summary", {}))

    # Section 8 — SEC Filing Review (Narrative Q&A)
    _heading(doc, "8.  SEC Filing Review — Qualitative Analysis")
    _note(doc, "The following findings are drawn from the company's SEC filings (10-K, 10-Q) using text analysis.")
    _narrative_qa_table(doc, state.get("narrative_analysis", []))

    # Section 9 — Fraud Detection & Financial Health Models
    _heading(doc, "9.  Fraud Detection & Financial Health Models")
    _note(doc, "These are standard quantitative models used by institutional investors to detect financial stress, earnings manipulation, and bankruptcy risk.")
    _fraud_detection_table(doc, state)

    # Chart: Risk Radar
    _embed_chart(doc, charts.get("risk_radar"), "Risk Model Radar — All 5 Models Normalised (10 = Safest)")
    # Chart: Risk Flowchart (matplotlib-drawn decision tree)
    _embed_chart(doc, charts.get("risk_flowchart"), "Risk Assessment Flowchart — 5 Models → Overall Verdict")


    # Section 10 — Valuation
    _heading(doc, "10. Valuation")
    _body(doc, structured.get("valuation_narrative", ""))
    _valuation_table(doc, state)

    # Chart: DCF Waterfall
    _embed_chart(doc, charts.get("dcf_waterfall"), "DCF Valuation Range — Monte Carlo Scenarios vs Current Market Price")
    # Chart: Valuation Flowchart (matplotlib-drawn decision tree)
    _embed_chart(doc, charts.get("valuation_flowchart"), "Valuation Decision Flowchart — DCF Scenarios vs Market Cap")

    # Chart: Balance Sheet Pie
    _embed_chart(doc, charts.get("balance_sheet_pie"), "Total Asset Composition — Where the Company's Money is Deployed")

    # Section 11 — Risk Flags
    flags = structured.get("risk_flags", [])
    if flags:
        _heading(doc, "11. Risk Flags")
        _risk_flags_table(doc, flags)

    # Section 12 — Data Unavailable
    _heading(doc, "12. Data Unavailable")
    _missing_data_section(doc, structured.get("missing_data", []))

    _disclaimer_page(doc, ticker)
    doc.save(str(out_path))
def _configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = DARK_RGB
def _cell_shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)
def _heading(doc: Document, title: str) -> None:
    """Gold background section heading with black text — high contrast."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(title.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DARK_RGB
    # Gold bottom border rule
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), XL_GOLD)
    pBdr.append(bottom)
    pPr.append(pBdr)
import re as _re
def _md_runs(para, text: str, base_size_pt: float = 10.5,
             base_color: RGBColor = None, italic: bool = False) -> None:
 
    if base_color is None:
        base_color = DARK_RGB
    pattern = _re.compile(r'(\*{3}.+?\*{3}|\*{2}.+?\*{2}|\*.+?\*)', _re.DOTALL)
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            r = para.add_run(text[last:m.start()])
            r.font.size = Pt(base_size_pt)
            r.font.color.rgb = base_color
            r.font.italic = italic

        raw = m.group(0)
        if raw.startswith('***'):
            inner = raw[3:-3]
            r = para.add_run(inner)
            r.font.bold = True
            r.font.italic = True
        elif raw.startswith('**'):
            inner = raw[2:-2]
            r = para.add_run(inner)
            r.font.bold = True
            r.font.italic = italic
        else:  # single *
            inner = raw[1:-1]
            r = para.add_run(inner)
            r.font.italic = True
            r.font.bold = False

        r.font.size = Pt(base_size_pt)
        r.font.color.rgb = base_color
        last = m.end()
    if last < len(text):
        r = para.add_run(text[last:])
        r.font.size = Pt(base_size_pt)
        r.font.color.rgb = base_color
        r.font.italic = italic
def _body(doc: Document, text: str) -> None:
    if not text:
        return
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()  # blank spacer
            continue
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        _md_runs(para, line, base_size_pt=10.5, base_color=DARK_RGB)
def _note(doc: Document, text: str) -> None:
    if not text:
        return
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    r = para.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
def _embed_chart(doc: Document, chart_path, caption: str = "") -> None:
    if not chart_path:
        return
    from pathlib import Path as _Path
    p = _Path(chart_path) if not isinstance(chart_path, _Path) else chart_path
    if not p.exists():
        return
    try:
        from docx.shared import Inches
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(8)
        spacer.paragraph_format.space_after = Pt(2)
        run = spacer.add_run()
        run.add_picture(str(p), width=Inches(6.0))
        spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Caption in italic grey below chart
        if caption:
            cap = doc.add_paragraph(f"Figure: {caption}")
            cap.runs[0].font.size = Pt(8.5)
            cap.runs[0].font.italic = True
            cap.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(10)
    except Exception as e:
        print(f"[embed_chart] Could not embed {p.name}: {e}")
def _tbl_header_cell(cell, text: str) -> None:
    """Navy background, white bold text."""
    _cell_shade(cell, XL_NAVY)
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = WHITE_RGB
    run.font.size = Pt(9.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
def _tbl_label_cell(cell, text: str, bold: bool = False, shade: str = XL_LIGHT_GOLD) -> None:
    _cell_shade(cell, shade)
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.bold = bold
    run.font.color.rgb = DARK_RGB
    run.font.size = Pt(9.5)
def _tbl_value_cell(cell, text: str, shade: str = XL_WHITE, color: RGBColor = None) -> None:
    _cell_shade(cell, shade)
    cell.text = str(text) if text is not None else "N/A"
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = color if color else DARK_RGB
    run.font.size = Pt(9.5)
def _add_title_page(doc: Document, ticker: str, structured: dict) -> None:
    if LOGO_PATH.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(str(LOGO_PATH), width=Inches(2.2))
    else:
        p = doc.add_paragraph("THE WALL STREET DD")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(20)
        p.runs[0].font.color.rgb = GOLD_RGB
    doc.add_paragraph()
    firm = doc.add_paragraph("The Wall Street DD")
    firm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firm.runs[0].font.size = Pt(13)
    firm.runs[0].font.bold = True
    firm.runs[0].font.color.rgb = GOLD_RGB
    sub = doc.add_paragraph("Autonomous Financial Due Diligence")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()
    doc.add_paragraph()
    t1 = doc.add_paragraph("FINANCIAL DUE DILIGENCE REPORT")
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1.runs[0].font.size = Pt(20)
    t1.runs[0].font.bold = True
    t1.runs[0].font.color.rgb = DARK_RGB
    t2 = doc.add_paragraph(ticker)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.size = Pt(36)
    t2.runs[0].font.bold = True
    t2.runs[0].font.color.rgb = GOLD_RGB
    doc.add_paragraph()
    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    for i, (lbl, val) in enumerate([
        ("Date of Report", datetime.now().strftime("%B %d, %Y")),
        ("Ticker Symbol", ticker),
        ("Classification", "CONFIDENTIAL — For Institutional Use Only"),
    ]):
        _cell_shade(meta.cell(i, 0), XL_NAVY)
        meta.cell(i, 0).text = lbl
        meta.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        meta.cell(i, 0).paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
        meta.cell(i, 0).paragraphs[0].runs[0].font.size = Pt(9)
        meta.cell(i, 1).text = val
        meta.cell(i, 1).paragraphs[0].runs[0].font.size = Pt(9)
        meta.cell(i, 1).paragraphs[0].runs[0].font.color.rgb = DARK_RGB
    meta.columns[0].width = Inches(2.0)
    meta.columns[1].width = Inches(4.0)
    doc.add_paragraph()
    conf = doc.add_paragraph(
        "CONFIDENTIAL. This report is prepared solely for the intended recipient. "
        "No part may be reproduced or distributed without prior written consent of The Wall Street DD."
    )
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf.runs[0].font.size = Pt(8)
    conf.runs[0].font.italic = True
    conf.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)
def _risk_banner(doc: Document, rating: str) -> None:
    r = rating.upper()
    if "LOW" in r or "HEALTHY" in r:
        bg, text_color = "1E8B4C", WHITE_RGB
    elif "HIGH" in r or "STRESS" in r:
        bg, text_color = "C0392B", WHITE_RGB
    else:
        bg, text_color = "E67E22", WHITE_RGB

    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _cell_shade(cell, bg)
    cell.text = f"  {rating}  "
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(14)
    cell.paragraphs[0].runs[0].font.color.rgb = text_color
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
def _financial_statements_table(doc: Document, state: Dict[str, Any]) -> None:
    rows_data = [
        ("Total Revenue",          "Money earned from selling products/services",                   "revenue_latest",                      "revenue_previous",                      False),
        ("Cost of Revenue",        "Direct cost to produce what was sold",                          "cogs_latest",                         "cogs_previous",                         False),
        ("Gross Profit",           "Revenue minus direct production costs",                         "gross_profit_latest",                 "gross_profit_previous",                 False),
        ("SG&A Expenses",          "Selling, admin & general overhead costs",                       "sga_expenses_latest",                 "sga_expenses_previous",                 False),
        ("Depreciation & Amort.", "Non-cash charge for asset wear-and-tear",                        "depreciation_amortization_latest",    "depreciation_amortization_previous",    False),
        ("Operating Income",       "Profit from core operations before interest/taxes",              "operating_income_latest",             "operating_income_previous",             False),
        ("Net Profit",             "Final profit after ALL costs including taxes",                   "net_income_latest",                   "net_income_previous",                   False),
        ("Income Tax Paid",        "Taxes paid to governments on profits",                           "income_tax_latest",                   "income_tax_previous",                   False),
        ("Interest on Debt",       "Cost paid to lenders on outstanding borrowings",                 "interest_expense_latest",             "interest_expense_previous",             False),
        ("Cash from Operations",   "Actual cash the business generated (not accounting profit)",     "operating_cash_flow_latest",          "operating_cash_flow_previous",          False),
        ("Capital Expenditure",    "Cash spent buying/upgrading physical assets",                    "capex_latest",                        "capex_previous",                        False),
        ("Cash & Equivalents",     "Immediately available cash and near-cash assets",                "cash_and_equivalents_latest",         "cash_and_equivalents_previous",         False),
        ("Money Owed to Company",  "Revenue earned but not yet collected from customers",            "receivables_latest",                  "receivables_previous",                  False),
        ("Unsold Inventory",       "Goods produced but not yet sold",                               "inventory_latest",                    "inventory_previous",                    False),
        ("Short-term Assets",      "Assets convertible to cash within 12 months",                   "current_assets_latest",               "current_assets_previous",               False),
        ("Short-term Obligations", "Bills and debts due within 12 months",                          "current_liabilities_latest",          "current_liabilities_previous",          False),
        ("Total Assets",           "Everything the company owns",                                   "total_assets_latest",                 "total_assets_previous",                 False),
        ("Total Obligations",      "Everything the company owes",                                   "total_liabilities_latest",            "total_liabilities_previous",            False),
        ("Long-term Debt",         "Loans and bonds due in more than 12 months",                    "long_term_debt_latest",               "long_term_debt_previous",               False),
        ("Short-term Debt",        "Borrowings due within 12 months",                               "short_term_debt_latest",              "short_term_debt_previous",              False),
        ("Shareholders Equity",    "What investors own = Total Assets minus Total Obligations",      "shareholders_equity_latest",          "shareholders_equity_previous",          False),
        ("Retained Profits",       "Cumulative profits kept inside the company over its history",    "retained_earnings_latest",            "retained_earnings_previous",            False),
    ]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    for i, hd in enumerate(["Metric", "What It Measures", "Latest Year", "Previous Year"]):
        _tbl_header_cell(tbl.rows[0].cells[i], hd)
    for idx, (label, desc, lk, pk, is_pct) in enumerate(rows_data):
        lv = state.get(lk)
        pv = state.get(pk)
        if lv is None and pv is None:
            continue
        row = tbl.add_row()
        shade = XL_LIGHT_GOLD if idx % 2 == 0 else XL_WHITE
        _tbl_label_cell(row.cells[0], label, bold=True, shade=shade)
        _tbl_label_cell(row.cells[1], desc, shade=shade)
        _tbl_value_cell(row.cells[2], _fmt(lv, is_pct), shade=shade)
        _tbl_value_cell(row.cells[3], _fmt(pv, is_pct), shade=shade)
    tbl.columns[0].width = Inches(1.5)
    tbl.columns[1].width = Inches(2.8)
    tbl.columns[2].width = Inches(1.35)
    tbl.columns[3].width = Inches(1.35)
    doc.add_paragraph()

def _kpi_table(doc: Document, state: Dict[str, Any]) -> None:
    quant = state.get("quant_analysis") or {}
    if not quant:
        _note(doc, "No quantitative scores available — scoring engine did not return results.")
        return
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    for i, hd in enumerate(["Metric", "What It Measures", "Value", "Health Score\n(1–10)", "Finding"]):
        _tbl_header_cell(tbl.rows[0].cells[i], hd)
    for idx, (metric_key, data) in enumerate(quant.items()):
        label_info = METRIC_LABELS.get(metric_key, (metric_key, ""))
        display_name = label_info[0]
        description  = label_info[1]

        raw_value = data.get("value")
        audit     = data.get("audit", {})
        score     = audit.get("health_score")
        assessment = audit.get("assessment", "")
        risk_flag = audit.get("risk_flag", False)

        is_pct = metric_key in ("GrossMargin", "OperatingMargin", "NetProfitMargin",
                                "ReturnOnAssets", "ReturnOnEquity")
        is_ratio = metric_key in ("CurrentRatio", "DebtToEquity")

        if is_pct and raw_value is not None:
            val_str = f"{raw_value * 100:.2f}%"
        elif is_ratio and raw_value is not None:
            val_str = f"{raw_value:.2f}x"
        else:
            val_str = _fmt_currency(raw_value) if raw_value is not None else "N/A"

        score_str = f"{score}/10" if score is not None else "N/A"

        row = tbl.add_row()
        shade = XL_LIGHT_GOLD if idx % 2 == 0 else XL_WHITE
        _tbl_label_cell(row.cells[0], display_name, bold=True, shade=shade)
        _tbl_label_cell(row.cells[1], description, shade=shade)
        _tbl_value_cell(row.cells[2], val_str, shade=shade)
        score_cell = row.cells[3]
        if score is not None:
            if score >= 8:
                _cell_shade(score_cell, "1E8B4C")
                score_cell.text = score_str
                score_cell.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
            elif score >= 5:
                _cell_shade(score_cell, "E67E22")
                score_cell.text = score_str
                score_cell.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
            else:
                _cell_shade(score_cell, "C0392B")
                score_cell.text = score_str
                score_cell.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
        else:
            _tbl_value_cell(score_cell, "N/A", shade=shade)
        score_cell.paragraphs[0].runs[0].font.bold = True
        score_cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        score_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        flag_prefix = "[RISK FLAG] " if risk_flag else ""
        _tbl_value_cell(row.cells[4], flag_prefix + assessment, shade=shade)
        if risk_flag:
            row.cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    tbl.columns[0].width = Inches(1.4)
    tbl.columns[1].width = Inches(1.9)
    tbl.columns[2].width = Inches(0.9)
    tbl.columns[3].width = Inches(0.85)
    tbl.columns[4].width = Inches(2.0)
    doc.add_paragraph()
def _risk_assessment_section(doc: Document, risk: dict) -> None:
    categories = [
        ("Severe Risks",   risk.get("severe_risks", []),   "C0392B", WHITE_RGB, "FDEDED"),
        ("Moderate Risks", risk.get("moderate_risks", []), "E67E22", WHITE_RGB, "FEF3E7"),
        ("Positive Factors / Low Risks", risk.get("low_risks", []), "1E8B4C", WHITE_RGB, "E8F8EE"),
    ]
    for cat_label, items, hdr_color, hdr_text, row_bg in categories:
        if not items:
            continue
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells[0]
        _cell_shade(hdr, hdr_color)
        hdr.text = cat_label.upper()
        hdr.paragraphs[0].runs[0].font.bold = True
        hdr.paragraphs[0].runs[0].font.color.rgb = hdr_text
        hdr.paragraphs[0].runs[0].font.size = Pt(10)
        for item in items:
            row = tbl.add_row()
            _cell_shade(row.cells[0], row_bg)
            row.cells[0].text = f"  {item}"
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = DARK_RGB
        doc.add_paragraph()
def _narrative_qa_table(doc: Document, narrative: Any) -> None:
    QUESTION_LABELS = [
   "Strategic Moat & Disruption: How does management identify its core competitive barriers and market entry protection?",
    "Strategic Moat & Disruption: What specific low-cost or technological disruptors does management note as direct competitive threats?",
    "Consumer Shift Rationale: What explanations does management offer for structural changes in consumer purchasing habits?",
    "Consumer Shift Rationale: How is a changing product mix or volume compression changing overall segment revenue?",
    "Pricing Power Execution: What is the specific execution mechanism for passing input cost inflation onto consumers via price hikes?",
    "Pricing Power Execution: How does management describe consumer volume elasticity or pushback to recent price adjustments?",
    "Underpenetrated Growth Vectors: What is the company's expansion roadmap for new geographic or vertical market segments?",
    "Underpenetrated Growth Vectors: What targeted value proposition is being marketed to unlock these new addressable demographics?",
    "M&A Integration Friction: What operational friction points, timeline delays, or system infrastructure issues occurred during recent integrations?",
    "M&A Integration Friction: What explicit corporate restructuring charges or workforce alignment friction points were recorded?",
    "FX Operational Adaptation: What structural supply chain movements or local sourcing initiatives are being executed to combat currency issues?",
    "FX Operational Adaptation: How is localized pricing or contract renegotiation being deployed to mitigate foreign currency volatility?",
    "Incentive Compensation Alignment: What specific performance metrics (like EPS, ROIC, or total shareholder return) govern executive annual cash bonuses?",
    "Incentive Compensation Alignment: How are long-term equity awards insulated by clawback policies or equity holding requirements?",
    "Human Capital & Labor Vitality: What material turnover trends, attrition rates, or specialized engineering/software talent recruitment risks are disclosed?",
    "Human Capital & Labor Vitality: What is the current operational health of baseline labor relations and domestic/international workforce stability?",
    "Key-Person & Succession Risk: To what extent is the company structurally dependent on the continuous service of its founder or Chief Executive Officer?",
    "Key-Person & Succession Risk: What structural gaps or strategic execution risks exist regarding formal executive succession planning?",
    "Board Oversight of Emerging Technology: How are the board committees explicitly organized to oversee data privacy policies and cybersecurity governance?",
    "Board Oversight of Emerging Technology: What operational oversight frameworks govern corporate artificial intelligence implementation and risk tracking?",
    "Related-Party Transaction Risks: What related-party transactions involve real estate leases, asset purchases, or vendor relationships owned by executives?",
    "Related-Party Transaction Risks: What tracking controls and internal evaluation policies govern the review of major shareholder transactions?",
    "ESG & Decarbonization Mandates: What fundamental capital allocation shifts are required to achieve targeted net-zero or carbon emission metrics?",
    "ESG & Decarbonization Mandates: What explicit material risks do physical climate disruptions or transitional carbon regulations pose to manufacturing operations?",
    "Single-Source Sourcing Chokepoints: Where does the business rely on a single or sole-source vendor for critical production components?",
    "Single-Source Sourcing Chokepoints: What internal raw material dependency risks lack alternative qualification paths or backup supply buffers?",
    "Logistics & Manufacturing Concentration: What geographical or physical concentration risks exist regarding the company's internal manufacturing facilities?",
    "Logistics & Manufacturing Concentration: What single-point distribution vulnerabilities exist across the outsourced third-party logistics network?",
    "Raw Material Scarcity Strategies: What are the operational parameters and minimum volume requirements of the company's long-term procurement obligations?",
    "Raw Material Scarcity Strategies: How has supplier concentration shifted structural input cost pricing power away from the company?",
]
    if not narrative:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        _tbl_header_cell(tbl.rows[0].cells[0], "Question")
        _tbl_header_cell(tbl.rows[0].cells[1], "Finding from SEC Filings")
        for q in QUESTION_LABELS:
            row = tbl.add_row()
            _tbl_label_cell(row.cells[0], q, shade=XL_LIGHT_GOLD)
            _cell_shade(row.cells[1], "FEF3E7")
            row.cells[1].text = "Not available — SEC filing database connection required to populate this section."
            row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        tbl.columns[0].width = Inches(2.5)
        tbl.columns[1].width = Inches(4.5)
        doc.add_paragraph()
        return

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    _tbl_header_cell(tbl.rows[0].cells[0], "Question")
    _tbl_header_cell(tbl.rows[0].cells[1], "Finding from SEC Filings")

    for idx, item in enumerate(narrative):
        if not isinstance(item, dict):
            continue
        for raw_q, answer in item.items():
            display_q = QUESTION_LABELS[idx] if idx < len(QUESTION_LABELS) else raw_q[:120]
            row = tbl.add_row()
            shade = XL_LIGHT_GOLD if idx % 2 == 0 else XL_WHITE
            _tbl_label_cell(row.cells[0], display_q, shade=shade)
            _tbl_value_cell(row.cells[1], str(answer)[:800], shade=shade)
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

    tbl.columns[0].width = Inches(2.5)
    tbl.columns[1].width = Inches(4.5)
    doc.add_paragraph()
def _fraud_detection_table(doc: Document, state: Dict[str, Any]) -> None:
    risk_report = state.get("risk_report", {})

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    for i, hd in enumerate(["Model", "What It Tests", "How to Read the Score", "Score", "Risk Level"]):
        _tbl_header_cell(tbl.rows[0].cells[i], hd)

    for idx, (key, meta) in enumerate(FRAUD_MODEL_LABELS.items()):
        score_val = state.get(meta["field"])
        rr = risk_report.get(key, {})
        flag = rr.get("risk", "N/A")

        if score_val is None:
            score_str = "N/A"
        elif key == "ohlson":
            score_str = f"{score_val * 100:.4f}% probability"
        elif key == "merton":
            score_str = f"{score_val:.3f} std deviations"
        elif key == "piotroski":
            score_str = f"{int(score_val)}/9"
        else:
            score_str = f"{score_val:.3f}"

        row = tbl.add_row()
        shade = XL_LIGHT_GOLD if idx % 2 == 0 else XL_WHITE
        _tbl_label_cell(row.cells[0], meta["name"], bold=True, shade=shade)
        _tbl_value_cell(row.cells[1], meta["what_it_is"], shade=shade)
        _tbl_value_cell(row.cells[2], meta["how_to_read"], shade=shade)
        _tbl_value_cell(row.cells[3], score_str, shade=shade)
        flag_cell = row.cells[4]
        flag_upper = flag.upper()
        if "HIGH" in flag_upper:
            _cell_shade(flag_cell, "C0392B")
            flag_cell.text = flag
            flag_cell.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
        elif "MEDIUM" in flag_upper:
            _cell_shade(flag_cell, "E67E22")
            flag_cell.text = flag
            flag_cell.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
        else:
            _cell_shade(flag_cell, "1E8B4C")
            flag_cell.text = flag
            flag_cell.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
        flag_cell.paragraphs[0].runs[0].font.bold = True
        flag_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    tbl.columns[0].width = Inches(1.5)
    tbl.columns[1].width = Inches(1.8)
    tbl.columns[2].width = Inches(1.8)
    tbl.columns[3].width = Inches(1.2)
    tbl.columns[4].width = Inches(0.9)
    doc.add_paragraph()
    dcf = risk_report.get("dcf_risk", {})
    if dcf:
        doc.add_paragraph(
            "Valuation Uncertainty (Monte Carlo Simulation — 10,000 scenarios run):"
        ).runs[0].font.bold = True
        vtbl = doc.add_table(rows=1, cols=4)
        vtbl.style = "Table Grid"
        for i, hd in enumerate(["Scenario", "What It Means", "Value", "Risk Level"]):
            _tbl_header_cell(vtbl.rows[0].cells[i], hd)

        scenarios = [
            ("Bear Case (P10 Floor)",    "Worst outcome in 1 in 10 simulations",  dcf.get("p10")),
            ("Base Case (P50 Median)",   "Most likely outcome across all simulations", dcf.get("p50")),
            ("Bull Case (P90 Ceiling)",  "Best outcome in 1 in 10 simulations",   dcf.get("p90")),
        ]
        dcf_flag = dcf.get("risk", "N/A")
        for i, (label, desc, val) in enumerate(scenarios):
            row = vtbl.add_row()
            shade = XL_LIGHT_GOLD if i % 2 == 0 else XL_WHITE
            _tbl_label_cell(row.cells[0], label, bold=True, shade=shade)
            _tbl_value_cell(row.cells[1], desc, shade=shade)
            _tbl_value_cell(row.cells[2], _fmt_currency(val), shade=shade)
            if i == 1:  
                fc = row.cells[3]
                if "HIGH" in dcf_flag.upper():
                    _cell_shade(fc, "C0392B"); fc.text = dcf_flag; fc.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
                elif "MEDIUM" in dcf_flag.upper():
                    _cell_shade(fc, "E67E22"); fc.text = dcf_flag; fc.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
                else:
                    _cell_shade(fc, "1E8B4C"); fc.text = dcf_flag; fc.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
                fc.paragraphs[0].runs[0].font.bold = True
                fc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                _tbl_value_cell(row.cells[3], "", shade=shade)
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

        vtbl.columns[0].width = Inches(1.5)
        vtbl.columns[1].width = Inches(2.5)
        vtbl.columns[2].width = Inches(1.5)
        vtbl.columns[3].width = Inches(1.5)
        doc.add_paragraph()
def _valuation_table(doc: Document, state: Dict[str, Any]) -> None:
    rows_data = [
        ("Intrinsic Value — Single Estimate (Deterministic DCF)",
         "A single calculated 'fair value' using one set of assumptions. Think of this as the analyst's base-case estimate.",
         state.get("deterministic_dcf_value")),
        ("Bear Case Value — P10 Floor (Monte Carlo)",
         "If the future turns out worse than expected, this is the minimum value in 90% of scenarios simulated.",
         state.get("monte_carlo_p10_floor")),
        ("Base Case Value — P50 Median (Monte Carlo)",
         "The most likely intrinsic value based on 10,000 simulated future scenarios.",
         state.get("monte_carlo_p50_median")),
        ("Bull Case Value — P90 Ceiling (Monte Carlo)",
         "If the future is better than expected, this is the maximum value in 90% of scenarios simulated.",
         state.get("monte_carlo_p90_ceiling")),
        ("Current Stock Price (Market Price)",
         "What the stock costs to buy today in the open market.",
         state.get("current_equity_price")),
        ("Market Capitalisation",
         "Total market value of all company shares. Stock Price × Total Shares Outstanding.",
         state.get("market_capitalization")),
        ("10-Year US Treasury Yield (Risk-Free Rate)",
         "The return an investor can get with zero risk (US government bonds). Used as the baseline in valuation.",
         state.get("risk_free_rate")),
        ("Annualised Stock Price Volatility (252 trading days)",
         "How much the stock price swings up and down over a year. Higher = more uncertainty.",
         state.get("historical_equity_volatility_252d")),
    ]

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for i, hd in enumerate(["Metric", "Plain-English Meaning", "Value"]):
        _tbl_header_cell(tbl.rows[0].cells[i], hd)

    for idx, (label, desc, val) in enumerate(rows_data):
        if val is None:
            continue
        row = tbl.add_row()
        shade = XL_LIGHT_GOLD if idx % 2 == 0 else XL_WHITE
        _tbl_label_cell(row.cells[0], label, bold=True, shade=shade)
        _tbl_value_cell(row.cells[1], desc, shade=shade)
        # Format: percentages vs currency vs raw
        if label in ("10-Year US Treasury Yield (Risk-Free Rate)",
                     "Annualised Stock Price Volatility (252 trading days)"):
            fmt_val = f"{float(val) * 100:.2f}%" if val < 2 else f"{float(val):.2f}%"
        elif label == "Current Stock Price (Market Price)":
            fmt_val = f"${float(val):.2f}"
        else:
            fmt_val = _fmt_currency(val)
        _tbl_value_cell(row.cells[2], fmt_val, shade=shade)
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    tbl.columns[0].width = Inches(2.2)
    tbl.columns[1].width = Inches(3.0)
    tbl.columns[2].width = Inches(1.8)
    doc.add_paragraph()
def _risk_flags_table(doc: Document, flags: list) -> None:
    for flag in flags:
        tbl = doc.add_table(rows=1, cols=1)
        cell = tbl.cell(0, 0)
        _cell_shade(cell, "FDEDED")
        cell.text = f"  {flag}"
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    doc.add_paragraph()
def _missing_data_section(doc: Document, missing: list) -> None:
    if not missing:
        _body(doc, "All required data fields were available for this analysis.")
        return
    for item in missing:
        p = doc.add_paragraph(style="List Bullet")
        _md_runs(p, str(item), base_size_pt=10)
    doc.add_paragraph()
def _disclaimer_page(doc: Document, ticker: str) -> None:
    doc.add_page_break()
    _heading(doc, "Disclaimer")
    _body(doc,
        f"This report has been generated by The Wall Street DD autonomous due-diligence platform "
        f"using publicly available financial data for {ticker}. It is intended for institutional "
        f"use only and does not constitute investment advice. Past performance is not indicative "
        f"of future results. The Wall Street DD makes no warranties regarding the accuracy or "
        f"completeness of this report. Recipients should conduct their own independent analysis "
        f"before making investment decisions."
    )
def _build_excel(state: Dict[str, Any], ticker: str, out_path: Path) -> None:
    wb = Workbook()
    wb.remove(wb.active)
    _sheet_dcf(wb, state, ticker)
    _sheet_risk_scores(wb, state)
    _sheet_financial_data(wb, state)
    _sheet_monte_carlo(wb, state)
    wb.save(str(out_path))
def _xf(bold=False, size=10, color=XL_DARK, italic=False) -> Font:
    return Font(name="Calibri", bold=bold, size=size, color=color, italic=italic)
def _xfill(hex_color: str) -> PatternFill:
    return PatternFill("solid", fgColor=hex_color)
def _xborder() -> Border:
    s = Side(style="thin", color="C9A84C")
    return Border(left=s, right=s, top=s, bottom=s)
def _xhdr(ws, row: int, col: int, value: str) -> None:
    cell = ws.cell(row=row, column=col, value=value)
    cell.font = _xf(bold=True, size=10, color=XL_WHITE)
    cell.fill = _xfill(XL_NAVY)
    cell.border = _xborder()
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
def _xlbl(ws, row: int, col: int, value: str, bold=False) -> None:
    cell = ws.cell(row=row, column=col, value=value)
    cell.font = _xf(bold=bold, size=10)
    cell.fill = _xfill(XL_LIGHT_GOLD)
    cell.border = _xborder()
    cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)


def _xval(ws, row: int, col: int, value=None, fmt: str = None, formula: str = None) -> None:
    cell = ws.cell(row=row, column=col, value=formula if formula else value)
    cell.font = _xf(size=10)
    cell.fill = _xfill(XL_WHITE)
    cell.border = _xborder()
    cell.alignment = Alignment(horizontal="right", vertical="center")
    if fmt:
        cell.number_format = fmt


def _xtitle(ws, row: int, title: str, span: int = 10) -> None:
    cell = ws.cell(row=row, column=1, value=title)
    cell.font = _xf(bold=True, size=14, color=XL_GOLD)
    cell.fill = _xfill(XL_NAVY)
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=span)
    cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[row].height = 30


def _xsection(ws, row: int, title: str, span: int = 10) -> None:
    cell = ws.cell(row=row, column=1, value=title)
    cell.font = _xf(bold=True, size=11, color=XL_SEC_TXT)
    cell.fill = _xfill(XL_SECTION)
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=span)
    cell.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[row].height = 22



def _sheet_dcf(wb: Workbook, state: Dict[str, Any], ticker: str) -> None:
    ws = wb.create_sheet("DCF Model")
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "C4"

    for i, w in enumerate([3, 42, 22, 44, 20, 20, 20, 20, 20, 20], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    _xtitle(ws, 1, f"THE WALL STREET DD  —  {ticker}  —  FULL DISCOUNTED CASH FLOW (DCF) MODEL", 10)
    ws.cell(row=2, column=1, value=f"Generated: {datetime.now().strftime('%B %d, %Y')}  |  CONFIDENTIAL").font = _xf(italic=True, size=9, color="808080")

    revenue     = state.get("revenue_latest") or 0.0
    net_income  = state.get("net_income_latest") or 0.0
    op_cf       = state.get("operating_cash_flow_latest") or 0.0
    capex       = abs(state.get("capex_latest") or 0.0)
    lt_debt     = state.get("long_term_debt_latest") or 0.0
    st_debt     = state.get("short_term_debt_latest") or 0.0
    total_debt  = lt_debt + st_debt
    equity      = state.get("shareholders_equity_latest") or 0.0
    risk_free   = state.get("risk_free_rate") or 0.04
    volatility  = state.get("historical_equity_volatility_252d") or 0.25
    market_cap  = state.get("market_capitalization") or 0.0
    price       = state.get("current_equity_price") or 0.0
    ebit_margin = (net_income / revenue) if revenue else 0.0

    _xsection(ws, 4, "A.  ACTUAL FINANCIAL INPUTS  (sourced from SEC EDGAR filings — do not edit)", 10)
    r = 5
    for col, hdr in enumerate(["Metric", "Value", "Plain-English Meaning"], 2):
        _xhdr(ws, r, col, hdr)
    r += 1

    input_rows = {}
    actual_inputs = [
        ("Total Revenue (Latest Year)",      revenue,       "All money earned from selling products/services",          '#,##0.00,,"M"'),
        ("EBIT Margin — Actual",             ebit_margin,   "Net Profit as % of Revenue (actual, not projected)",       '0.00%'),
        ("Operating Cash Flow",              op_cf,         "Actual cash generated from core operations",               '#,##0.00,,"M"'),
        ("Capital Expenditure (CapEx)",      capex,         "Cash spent buying/upgrading physical assets",              '#,##0.00,,"M"'),
        ("Free Cash Flow",                   op_cf - capex, "Cash left over after CapEx (OCF - CapEx)",                 '#,##0.00,,"M"'),
        ("Long-term Debt",                   lt_debt,       "Loans & bonds due in more than 12 months",                 '#,##0.00,,"M"'),
        ("Short-term Debt",                  st_debt,       "Borrowings due within 12 months",                          '#,##0.00,,"M"'),
        ("Total Debt",                       total_debt,    "All borrowings combined",                                  '#,##0.00,,"M"'),
        ("Shareholders Equity",              equity,        "What owners possess: Total Assets minus Total Liabilities", '#,##0.00,,"M"'),
        ("Market Capitalisation",            market_cap,    "Stock Price × Total Shares Outstanding",                   '#,##0.00,,"M"'),
        ("Current Stock Price",              price,         "Today's share price in the open market",                   '$#,##0.00'),
        ("10-Year Treasury Yield (Risk-Free Rate)", risk_free, "Baseline safe return — US government bonds",           '0.00%'),
        ("Annualised Stock Volatility (252 days)", volatility, "How much the stock price swings in a year",            '0.00%'),
    ]
    for label, value, note, fmt in actual_inputs:
        _xlbl(ws, r, 2, label)
        _xval(ws, r, 3, value, fmt=fmt)
        ws.cell(row=r, column=4, value=note).font = _xf(italic=True, size=9, color="808080")
        input_rows[label] = r
        r += 1

    r += 1
    _xsection(ws, r, "B.  EDITABLE ASSUMPTIONS  (highlighted cells — change these to stress-test the model)", 10)
    r += 1
    for col, hdr in enumerate(["Assumption", "Value", "Plain-English Meaning"], 2):
        _xhdr(ws, r, col, hdr)
    r += 1

    assumption_rows = {}
    assumptions = [
        ("Revenue Growth Rate — Years 1-5",         0.08,  "How fast revenue grows per year for first 5 years (8% = moderate growth)"),
        ("Revenue Growth Rate — Years 6-10",        0.05,  "How fast revenue grows per year for years 6–10 (5% = maturing growth)"),
        ("Terminal Growth Rate (After Year 10)",    0.025, "Long-run growth rate assumed forever after year 10 (matches long-run GDP)"),
        ("EBIT Margin — Projected",                 0.20,  "Operating profit as % of revenue assumed for future years"),
        ("Corporate Tax Rate",                      0.21,  "Percentage of operating profit paid as tax (US federal rate = 21%)"),
        ("Depreciation as % of Revenue",            0.03,  "Non-cash accounting charge estimated as % of revenue"),
        ("CapEx as % of Revenue",                   0.04,  "Investment in physical assets estimated as % of revenue"),
        ("Working Capital Change as % of Revenue",  0.01,  "Net change in short-term assets/liabilities as % of revenue"),
        ("Equity Risk Premium",                     0.055, "Extra return investors demand for owning stocks vs safe government bonds"),
        ("Beta",                                    1.0,   "How much the stock moves relative to the market (1.0 = moves with market)"),
        ("Cost of Debt (Pre-tax)",                  0.045, "Interest rate the company pays on its borrowings"),
        ("Debt Weight in Capital Structure",        0.30,  "What % of the company's funding comes from debt vs equity"),
        ("Projection Period (Years)",               10,    "How many years of cash flows we explicitly forecast before terminal value"),
    ]
    for label, value, note in assumptions:
        ws.cell(row=r, column=2, value=label).font = _xf(size=10)
        ws.cell(row=r, column=2).fill = _xfill(XL_LIGHT_GOLD)
        ws.cell(row=r, column=2).border = _xborder()
        val_cell = ws.cell(row=r, column=3, value=value)
        val_cell.font = _xf(bold=True, size=10, color="1A1A00")
        val_cell.fill = PatternFill("solid", fgColor="FFF9C4")  # Bright yellow = user-editable
        val_cell.border = _xborder()
        val_cell.alignment = Alignment(horizontal="right")
        val_cell.number_format = '0.00x' if label == "Beta" else ('0.00%' if isinstance(value, float) else '0')
        ws.cell(row=r, column=4, value=note).font = _xf(italic=True, size=9, color="808080")
        assumption_rows[label] = r
        r += 1

    # Revenue projection rows with real Excel formulas
    r += 2
    _xsection(ws, r, "C.  10-YEAR REVENUE & FREE CASH FLOW PROJECTIONS  (auto-calculated from assumptions above)", 10)
    r += 1
    yr_headers = ["Metric"] + [f"Year {y}" for y in range(1, 11)]
    for col, hdr in enumerate(yr_headers, 2):
        _xhdr(ws, r, col, hdr)
    r += 1

    g1 = assumption_rows["Revenue Growth Rate — Years 1-5"]
    g2 = assumption_rows["Revenue Growth Rate — Years 6-10"]
    rev_r = r
    ws.cell(row=r, column=2, value="Revenue ($M)").font = _xf(bold=True)
    ws.cell(row=r, column=2).fill = _xfill(XL_LIGHT_GOLD)
    ws.cell(row=r, column=2).border = _xborder()
    base_rev_row = input_rows["Total Revenue (Latest Year)"]
    for yi in range(10):
        col = 3 + yi
        if yi == 0:
            formula = f"=C{base_rev_row}*(1+C{g1})"
        elif yi < 5:
            formula = f"={get_column_letter(col-1)}{rev_r}*(1+C{g1})"
        else:
            formula = f"={get_column_letter(col-1)}{rev_r}*(1+C{g2})"
        c = ws.cell(row=r, column=col, value=formula)
        c.font = _xf(size=10); c.fill = _xfill(XL_WHITE); c.border = _xborder()
        c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    r += 1

    ebit_r = r
    em_row = assumption_rows["EBIT Margin — Projected"]
    ws.cell(row=r, column=2, value="EBIT — Operating Profit ($M)").font = _xf(size=10)
    ws.cell(row=r, column=2).fill = _xfill(XL_LIGHT_GOLD); ws.cell(row=r, column=2).border = _xborder()
    for yi in range(10):
        col = 3 + yi
        c = ws.cell(row=r, column=col, value=f"={get_column_letter(col)}{rev_r}*C{em_row}")
        c.font = _xf(size=10); c.fill = _xfill(XL_WHITE); c.border = _xborder()
        c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    r += 1

    tax_row = assumption_rows["Corporate Tax Rate"]
    nopat_r = r
    ws.cell(row=r, column=2, value="NOPAT — Net Operating Profit After Tax ($M)").font = _xf(size=10)
    ws.cell(row=r, column=2).fill = _xfill(XL_LIGHT_GOLD); ws.cell(row=r, column=2).border = _xborder()
    for yi in range(10):
        col = 3 + yi
        c = ws.cell(row=r, column=col, value=f"={get_column_letter(col)}{ebit_r}*(1-C{tax_row})")
        c.font = _xf(size=10); c.fill = _xfill(XL_WHITE); c.border = _xborder()
        c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    r += 1

    dna_row_assum = assumption_rows["Depreciation as % of Revenue"]
    dna_r = r
    ws.cell(row=r, column=2, value="Depreciation & Amortization ($M)").font = _xf(size=10)
    ws.cell(row=r, column=2).fill = _xfill(XL_LIGHT_GOLD); ws.cell(row=r, column=2).border = _xborder()
    for yi in range(10):
        col = 3 + yi
        c = ws.cell(row=r, column=col, value=f"={get_column_letter(col)}{rev_r}*C{dna_row_assum}")
        c.font = _xf(size=10); c.fill = _xfill(XL_WHITE); c.border = _xborder()
        c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    r += 1

    capex_assum = assumption_rows["CapEx as % of Revenue"]
    capex_r = r
    ws.cell(row=r, column=2, value="Capital Expenditure — CapEx ($M)").font = _xf(size=10)
    ws.cell(row=r, column=2).fill = _xfill(XL_LIGHT_GOLD); ws.cell(row=r, column=2).border = _xborder()
    for yi in range(10):
        col = 3 + yi
        c = ws.cell(row=r, column=col, value=f"={get_column_letter(col)}{rev_r}*C{capex_assum}")
        c.font = _xf(size=10); c.fill = _xfill(XL_WHITE); c.border = _xborder()
        c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    r += 1

    nwc_assum = assumption_rows["Working Capital Change as % of Revenue"]
    nwc_r = r
    ws.cell(row=r, column=2, value="Change in Working Capital ($M)").font = _xf(size=10)
    ws.cell(row=r, column=2).fill = _xfill(XL_LIGHT_GOLD); ws.cell(row=r, column=2).border = _xborder()
    for yi in range(10):
        col = 3 + yi
        c = ws.cell(row=r, column=col, value=f"={get_column_letter(col)}{rev_r}*C{nwc_assum}")
        c.font = _xf(size=10); c.fill = _xfill(XL_WHITE); c.border = _xborder()
        c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    r += 1

    fcf_r = r
    ws.cell(row=r, column=2, value="FREE CASH FLOW ($M)  =  NOPAT + D&A − CapEx − ΔWCC").font = _xf(bold=True, size=10, color=XL_GOLD)
    ws.cell(row=r, column=2).fill = _xfill(XL_NAVY); ws.cell(row=r, column=2).border = _xborder()
    for yi in range(10):
        col = 3 + yi
        lc = get_column_letter(col)
        c = ws.cell(row=r, column=col, value=f"={lc}{nopat_r}+{lc}{dna_r}-{lc}{capex_r}-{lc}{nwc_r}")
        c.font = _xf(bold=True, size=10, color=XL_GOLD); c.fill = _xfill(XL_NAVY); c.border = _xborder()
        c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    r += 1

    # WACC
    r += 2
    _xsection(ws, r, "D.  WEIGHTED AVERAGE COST OF CAPITAL (WACC)  —  The discount rate applied to future cash flows", 10)
    r += 1
    erp_row  = assumption_rows["Equity Risk Premium"]
    beta_row = assumption_rows["Beta"]
    rf_row   = input_rows["10-Year Treasury Yield (Risk-Free Rate)"]
    cod_row  = assumption_rows["Cost of Debt (Pre-tax)"]
    wd_row   = assumption_rows["Debt Weight in Capital Structure"]

    wacc_calcs = [
        ("Cost of Equity  =  Risk-Free Rate  +  Beta × Equity Risk Premium",
         f"=C{rf_row}+C{beta_row}*C{erp_row}", "CAPM formula — minimum return equity investors require"),
        ("After-tax Cost of Debt  =  Cost of Debt × (1 − Tax Rate)",
         f"=C{cod_row}*(1-C{tax_row})", "Debt is cheaper than equity because interest is tax-deductible"),
        ("Weight of Equity in Capital Structure",
         f"=1-C{wd_row}", "Proportion of the company's financing that comes from equity"),
        ("Weight of Debt in Capital Structure",
         f"=C{wd_row}", "Proportion of the company's financing that comes from debt"),
    ]
    ke_row = r
    for label, formula, note in wacc_calcs:
        _xlbl(ws, r, 2, label, bold=True)
        c = ws.cell(row=r, column=3, value=formula)
        c.font = _xf(size=10); c.fill = _xfill(XL_WHITE); c.border = _xborder()
        c.number_format = "0.00%"; c.alignment = Alignment(horizontal="right")
        ws.cell(row=r, column=4, value=note).font = _xf(italic=True, size=9, color="808080")
        r += 1

    wacc_r = r
    _xlbl(ws, r, 2, "WACC  (Weighted Average Cost of Capital)  —  the rate used to discount future cash flows", bold=True)
    wc = ws.cell(row=r, column=3, value=f"=C{ke_row+2}*C{ke_row}+C{ke_row+3}*C{ke_row+1}")
    wc.font = _xf(bold=True, size=12, color=XL_GOLD); wc.fill = _xfill(XL_NAVY); wc.border = _xborder()
    wc.number_format = "0.00%"; wc.alignment = Alignment(horizontal="right")
    ws.cell(row=r, column=4, value="Equity Weight × Cost of Equity  +  Debt Weight × After-tax Cost of Debt").font = _xf(italic=True, size=9, color="808080")
    r += 1

    # DCF Valuation
    r += 2
    _xsection(ws, r, "E.  DCF VALUATION  —  Present Value of All Future Cash Flows", 10)
    r += 1
    tgr_row  = assumption_rows["Terminal Growth Rate (After Year 10)"]
    proj_row = assumption_rows["Projection Period (Years)"]

    pv_r = r
    _xlbl(ws, r, 2, "Present Value of Each Year's Free Cash Flow ($M)", bold=True)
    ws.cell(row=r, column=2).fill = _xfill(XL_LIGHT_GOLD); ws.cell(row=r, column=2).border = _xborder()
    for yi, year in enumerate(range(1, 11)):
        col = 3 + yi
        lc = get_column_letter(col)
        c = ws.cell(row=r, column=col, value=f"={lc}{fcf_r}/((1+C{wacc_r})^{year})")
        c.font = _xf(size=10); c.fill = _xfill(XL_WHITE); c.border = _xborder()
        c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    r += 1

    sum_pv_r = r
    _xlbl(ws, r, 2, "Sum of All Present Values — Total PV of 10-year FCFs ($M)", bold=True)
    c = ws.cell(row=r, column=3, value=f"=SUM(C{pv_r}:{get_column_letter(12)}{pv_r})")
    c.font = _xf(bold=True); c.fill = _xfill(XL_WHITE); c.border = _xborder()
    c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    r += 1

    tv_r = r
    _xlbl(ws, r, 2, "Terminal Value — Value of ALL Cash Flows After Year 10 ($M)", bold=True)
    c = ws.cell(row=r, column=3,
                value=f"=({get_column_letter(12)}{fcf_r}*(1+C{tgr_row}))/(C{wacc_r}-C{tgr_row})")
    c.font = _xf(bold=True); c.fill = _xfill(XL_WHITE); c.border = _xborder()
    c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    ws.cell(row=r, column=4, value="Year 10 FCF × (1 + Terminal Growth Rate)  ÷  (WACC − Terminal Growth Rate)").font = _xf(italic=True, size=9, color="808080")
    r += 1

    pv_tv_r = r
    _xlbl(ws, r, 2, "Present Value of Terminal Value ($M)", bold=True)
    c = ws.cell(row=r, column=3, value=f"=C{tv_r}/((1+C{wacc_r})^C{proj_row})")
    c.font = _xf(bold=True); c.fill = _xfill(XL_WHITE); c.border = _xborder()
    c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    r += 1

    ev_r = r
    _xlbl(ws, r, 2, "ENTERPRISE VALUE  =  Sum of PV FCFs  +  PV Terminal Value ($M)", bold=True)
    c = ws.cell(row=r, column=3, value=f"=C{sum_pv_r}+C{pv_tv_r}")
    c.font = _xf(bold=True, size=12, color=XL_GOLD); c.fill = _xfill(XL_NAVY); c.border = _xborder()
    c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    r += 1

    cash = state.get("cash_and_equivalents_latest") or 0.0
    net_debt = total_debt - cash
    eq_r = r
    _xlbl(ws, r, 2, "EQUITY VALUE  =  Enterprise Value  −  Net Debt ($M)", bold=True)
    nd_cell = ws.cell(row=r, column=4, value=net_debt)
    nd_cell.number_format = '#,##0,,"M"'
    c = ws.cell(row=r, column=3, value=f"=C{ev_r}-D{r}")
    c.font = _xf(bold=True, size=12, color=XL_GOLD); c.fill = _xfill(XL_NAVY); c.border = _xborder()
    c.number_format = '#,##0,,"M"'; c.alignment = Alignment(horizontal="right")
    r += 1

    shares = state.get("common_shares_outstanding_latest")
    if shares and shares > 1000:
        _xlbl(ws, r, 2, "IMPLIED SHARE PRICE  =  Equity Value  ÷  Shares Outstanding", bold=True)
        sh_cell = ws.cell(row=r, column=4, value=shares)
        sh_cell.number_format = '#,##0'
        c = ws.cell(row=r, column=3, value=f"=C{eq_r}/D{r}")
        c.font = _xf(bold=True, size=14, color=XL_GOLD); c.fill = _xfill(XL_NAVY); c.border = _xborder()
        c.number_format = '$#,##0.00'; c.alignment = Alignment(horizontal="right")
    else:
        _xlbl(ws, r, 2, "Shares Outstanding — enter manually to calculate implied share price", bold=False)
        ws.cell(row=r, column=3, value="Enter shares outstanding in D column →").font = _xf(italic=True, size=9, color="808080")
        ws.cell(row=r, column=4, value=None).number_format = '#,##0'
    r += 1

    # Pipeline model output comparison
    r += 2
    _xsection(ws, r, "F.  PIPELINE MODEL OUTPUT  —  Scores computed by the analysis engine", 10)
    r += 1
    comparisons = [
        ("Deterministic DCF Value (Pipeline Engine)", state.get("deterministic_dcf_value") or 0,  '#,##0,,"B"'),
        ("Monte Carlo P10 — Bear Case Floor",         state.get("monte_carlo_p10_floor") or 0,    '#,##0,,"B"'),
        ("Monte Carlo P50 — Base Case Median",        state.get("monte_carlo_p50_median") or 0,   '#,##0,,"B"'),
        ("Monte Carlo P90 — Bull Case Ceiling",       state.get("monte_carlo_p90_ceiling") or 0,  '#,##0,,"B"'),
        ("Current Market Price (per share)",          state.get("current_equity_price") or 0,     '$#,##0.00'),
        ("Market Capitalisation",                     state.get("market_capitalization") or 0,    '#,##0,,"B"'),
    ]
    for label, val, fmt in comparisons:
        _xlbl(ws, r, 2, label)
        _xval(ws, r, 3, val, fmt=fmt)
        r += 1

    # Sensitivity table
    r += 2
    _xsection(ws, r, "G.  SENSITIVITY TABLE  —  Enterprise Value ($B) at Different WACC and Terminal Growth Rates", 10)
    r += 1
    wacc_range = [0.06, 0.07, 0.08, 0.09, 0.10, 0.11, 0.12]
    tgr_range  = [0.015, 0.020, 0.025, 0.030, 0.035]
    ws.cell(row=r, column=2, value="WACC → / Growth Rate ↓").font = _xf(bold=True, size=9)
    ws.cell(row=r, column=2).fill = _xfill(XL_NAVY); ws.cell(row=r, column=2).border = _xborder()
    ws.cell(row=r, column=2).font = _xf(bold=True, size=9, color=XL_WHITE)
    for wi, w in enumerate(wacc_range):
        c = ws.cell(row=r, column=3 + wi, value=f"{w*100:.0f}%")
        c.font = _xf(bold=True, size=9, color=XL_WHITE); c.fill = _xfill(XL_NAVY); c.border = _xborder()
        c.alignment = Alignment(horizontal="center")
    r += 1

    fcf_base = op_cf - capex
    for tgr in tgr_range:
        c = ws.cell(row=r, column=2, value=f"TGR {tgr*100:.1f}%")
        c.font = _xf(bold=True, size=9, color=XL_WHITE); c.fill = _xfill(XL_NAVY); c.border = _xborder()
        c.alignment = Alignment(horizontal="center")
        for wi, wacc in enumerate(wacc_range):
            try:
                g = 0.08
                fcf_y10 = fcf_base * ((1 + g) ** 10)
                tv_val = fcf_y10 * (1 + tgr) / (wacc - tgr) if wacc > tgr else None
                pv_tv_val = tv_val / ((1 + wacc) ** 10) if tv_val else None
                sum_pv_fcf = sum(fcf_base * ((1 + g) ** y) / ((1 + wacc) ** y) for y in range(1, 11))
                ev_val = (sum_pv_fcf + pv_tv_val) / 1e9 if pv_tv_val else None
            except Exception:
                ev_val = None
            cell = ws.cell(row=r, column=3 + wi, value=ev_val)
            if ev_val is not None and ev_val > 0:
                cell.number_format = '$#,##0.0"B"'
                cell.font = _xf(size=9)
                cell.fill = _xfill("E8F8EE")
            else:
                cell.value = "N/A"; cell.font = _xf(size=9, color=XL_RED)
            cell.border = _xborder()
            cell.alignment = Alignment(horizontal="right")
        r += 1


# ── Sheet 2: Risk Scores ─────────────────────────────────────────────────────

def _sheet_risk_scores(wb: Workbook, state: Dict[str, Any]) -> None:
    ws = wb.create_sheet("Risk Scores")
    ws.sheet_view.showGridLines = False
    for i, w in enumerate([2, 32, 18, 18, 45], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    _xtitle(ws, 1, "FRAUD DETECTION & FINANCIAL HEALTH MODELS — RISK SCORE ANALYSIS")
    r = 3
    _xsection(ws, r, "QUANTITATIVE RISK MODELS  —  Industry-standard mathematical models used by institutional investors")
    r += 1
    for col, hdr in enumerate(["Model", "Score", "Risk Level", "Interpretation"], 2):
        _xhdr(ws, r, col, hdr)
    r += 1

    for idx, (key, meta) in enumerate(FRAUD_MODEL_LABELS.items()):
        score_val = state.get(meta["field"])
        rr = state.get("risk_report", {}).get(key, {})
        flag = rr.get("risk", "N/A")

        if score_val is None:
            score_str = "N/A"
        elif key == "ohlson":
            score_str = f"{score_val * 100:.4f}% probability"
        elif key == "merton":
            score_str = f"{score_val:.3f} std deviations from default"
        elif key == "piotroski":
            score_str = f"{int(score_val)}/9"
        else:
            score_str = f"{score_val:.3f}"

        _xlbl(ws, r, 2, f"{meta['name']}\n{meta['what_it_is']}", bold=True)
        ws.row_dimensions[r].height = 40
        sc = ws.cell(row=r, column=3, value=score_str)
        sc.font = _xf(bold=True, size=11); sc.border = _xborder()
        sc.alignment = Alignment(horizontal="center", vertical="center")
        sc.fill = _xfill(XL_LIGHT_GOLD if idx % 2 == 0 else XL_WHITE)

        fc = ws.cell(row=r, column=4, value=flag)
        fc.font = _xf(bold=True, size=10, color=XL_WHITE)
        fc.border = _xborder()
        fc.alignment = Alignment(horizontal="center", vertical="center")
        if "HIGH" in flag.upper():
            fc.fill = _xfill(XL_RED)
        elif "MEDIUM" in flag.upper():
            fc.fill = _xfill(XL_AMBER)
        else:
            fc.fill = _xfill(XL_GREEN)

        ic = ws.cell(row=r, column=5, value=meta["how_to_read"])
        ic.font = _xf(size=9, italic=True); ic.border = _xborder()
        ic.fill = _xfill(XL_LIGHT_GOLD if idx % 2 == 0 else XL_WHITE)
        ic.alignment = Alignment(wrap_text=True, vertical="center")
        r += 1

    r += 2
    _xsection(ws, r, "OVERALL VERDICT")
    r += 1
    overall = state.get("risk_report", {}).get("overall_assessment", {})
    _xlbl(ws, r, 2, "Final Risk Verdict")
    vc = ws.cell(row=r, column=3, value=overall.get("final_verdict", "N/A"))
    vc.font = _xf(bold=True, size=12, color=XL_GOLD); vc.fill = _xfill(XL_NAVY); vc.border = _xborder()
    r += 1
    _xlbl(ws, r, 2, "Average Risk Score (1 = Low, 3 = High)")
    ac = ws.cell(row=r, column=3, value=overall.get("average_risk_score"))
    ac.font = _xf(bold=True, size=11); ac.fill = _xfill(XL_WHITE); ac.border = _xborder()
    ac.number_format = "0.00"


# ── Sheet 3: Financial Data ──────────────────────────────────────────────────

def _sheet_financial_data(wb: Workbook, state: Dict[str, Any]) -> None:
    ws = wb.create_sheet("Financial Data")
    ws.sheet_view.showGridLines = False
    for i, w in enumerate([2, 35, 22, 22], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    _xtitle(ws, 1, "EXTRACTED FINANCIAL DATA  —  Sourced from SEC EDGAR Filings")
    r = 3

    sections = [
        ("INCOME STATEMENT  —  How Much the Company Earned and Spent", [
            ("Total Revenue",               "revenue_latest",                       "revenue_previous"),
            ("Cost of Revenue",             "cogs_latest",                          "cogs_previous"),
            ("Gross Profit",                "gross_profit_latest",                  "gross_profit_previous"),
            ("SG&A Expenses",               "sga_expenses_latest",                  "sga_expenses_previous"),
            ("Depreciation & Amortization", "depreciation_amortization_latest",     "depreciation_amortization_previous"),
            ("Operating Income",            "operating_income_latest",              "operating_income_previous"),
            ("Net Profit",                  "net_income_latest",                    "net_income_previous"),
            ("Income Tax Paid",             "income_tax_latest",                    "income_tax_previous"),
            ("Interest on Debt",            "interest_expense_latest",              "interest_expense_previous"),
        ]),
        ("CASH FLOW STATEMENT  —  Actual Cash Moving In and Out", [
            ("Cash from Operations",        "operating_cash_flow_latest",           "operating_cash_flow_previous"),
            ("Capital Expenditure",         "capex_latest",                         "capex_previous"),
            ("Free Cash Flow (Derived)",    None,                                   None),
        ]),
        ("BALANCE SHEET  —  What the Company Owns and Owes", [
            ("Cash & Equivalents",          "cash_and_equivalents_latest",          "cash_and_equivalents_previous"),
            ("Money Owed to Company",       "receivables_latest",                   "receivables_previous"),
            ("Unsold Inventory",            "inventory_latest",                     "inventory_previous"),
            ("Short-term Assets Total",     "current_assets_latest",                "current_assets_previous"),
            ("Short-term Obligations",      "current_liabilities_latest",           "current_liabilities_previous"),
            ("Property, Plant & Equipment", "gross_ppe_latest",                     "gross_ppe_previous"),
            ("Total Assets",                "total_assets_latest",                  "total_assets_previous"),
            ("Total Obligations",           "total_liabilities_latest",             "total_liabilities_previous"),
            ("Long-term Debt",              "long_term_debt_latest",                "long_term_debt_previous"),
            ("Short-term Debt",             "short_term_debt_latest",               "short_term_debt_previous"),
            ("Shareholders Equity",         "shareholders_equity_latest",           "shareholders_equity_previous"),
            ("Retained Profits",            "retained_earnings_latest",             "retained_earnings_previous"),
        ]),
        ("MARKET DATA  —  Live Market Information", [
            ("Current Stock Price",         "current_equity_price",                 None),
            ("Market Capitalisation",       "market_capitalization",                None),
            ("Stock Volatility (252-day)",  "historical_equity_volatility_252d",    None),
            ("Risk-Free Rate",              "risk_free_rate",                       None),
            ("GNP Deflator",               "gnp_deflator",                         None),
        ]),
    ]

    for section_label, fields in sections:
        _xsection(ws, r, section_label)
        r += 1
        for col, hdr in enumerate(["Metric", "Latest Period", "Previous Period"], 2):
            _xhdr(ws, r, col, hdr)
        r += 1
        for idx, (label, lk, pk) in enumerate(fields):
            if label == "Free Cash Flow (Derived)" and lk is None:
                lv = _safe_float(state.get("operating_cash_flow_latest")) + _safe_float(state.get("capex_latest"))
                pv = _safe_float(state.get("operating_cash_flow_previous")) + _safe_float(state.get("capex_previous"))
            else:
                lv = state.get(lk) if lk else None
                pv = state.get(pk) if pk else None
            shade = XL_LIGHT_GOLD if idx % 2 == 0 else XL_WHITE
            _xlbl(ws, r, 2, label)
            ws.cell(row=r, column=2).fill = _xfill(shade)
            for col, val in [(3, lv), (4, pv)]:
                c = ws.cell(row=r, column=col, value=val)
                c.font = _xf(size=10); c.fill = _xfill(shade); c.border = _xborder()
                c.number_format = '#,##0,,"M"' if isinstance(val, float) and abs(val or 0) > 1_000_000 else '0.0000'
                c.alignment = Alignment(horizontal="right")
            r += 1
        r += 1


# ── Sheet 4: Monte Carlo ─────────────────────────────────────────────────────

def _sheet_monte_carlo(wb: Workbook, state: Dict[str, Any]) -> None:
    ws = wb.create_sheet("Monte Carlo")
    ws.sheet_view.showGridLines = False
    for i, w in enumerate([2, 32, 22, 45], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    _xtitle(ws, 1, "MONTE CARLO DCF SIMULATION  —  Valuation Range Under Uncertainty (10,000 Simulated Scenarios)")
    r = 3
    _xsection(ws, r, "SIMULATION OUTPUT  —  What different futures look like for this company's value")
    r += 1
    for col, hdr in enumerate(["Scenario", "Value", "Plain-English Meaning"], 2):
        _xhdr(ws, r, col, hdr)
    r += 1

    scenarios = [
        ("Bear Case — P10 Floor\n(Worst 10% of scenarios)",
         state.get("monte_carlo_p10_floor"),
         "If everything goes somewhat wrong — slower growth, margin pressure — the company is worth at least this much in 9 out of 10 simulated futures."),
        ("Base Case — P50 Median\n(Most likely outcome)",
         state.get("monte_carlo_p50_median"),
         "The middle value: half of simulated futures produce a higher value, half produce a lower value. This is the best single estimate."),
        ("Bull Case — P90 Ceiling\n(Best 10% of scenarios)",
         state.get("monte_carlo_p90_ceiling"),
         "If conditions are favourable — stronger growth, expanding margins — the company could be worth this much in roughly 1 in 10 futures."),
        ("Single-Point Estimate (Deterministic DCF)",
         state.get("deterministic_dcf_value"),
         "A single 'fair value' calculated with fixed base-case assumptions, without any randomness or scenario modelling."),
    ]

    fills = ["E8F8EE", XL_LIGHT_GOLD, "FEF3E7", XL_WHITE]
    for i, (label, val, interp) in enumerate(scenarios):
        _xlbl(ws, r, 2, label, bold=True)
        ws.cell(row=r, column=2).fill = _xfill(fills[i])
        ws.row_dimensions[r].height = 38
        c = ws.cell(row=r, column=3, value=val)
        c.font = _xf(bold=True, size=12); c.fill = _xfill(fills[i]); c.border = _xborder()
        c.number_format = '$#,##0,,"B"'; c.alignment = Alignment(horizontal="right", vertical="center")
        ic = ws.cell(row=r, column=4, value=interp)
        ic.font = _xf(size=9, italic=True); ic.fill = _xfill(fills[i]); ic.border = _xborder()
        ic.alignment = Alignment(wrap_text=True, vertical="center")
        r += 1

    r += 2
    _xsection(ws, r, "UPSIDE / DOWNSIDE ANALYSIS  —  How wide is the range of possible outcomes?")
    r += 1
    p10 = state.get("monte_carlo_p10_floor") or 0
    p50 = state.get("monte_carlo_p50_median") or 1
    p90 = state.get("monte_carlo_p90_ceiling") or 0
    upside   = (p90 - p50) / abs(p50) if p50 else 0
    downside = (p50 - p10) / abs(p50) if p50 else 0
    metrics = [
        ("Upside Potential  (Bull Case vs Base Case)", upside,         "0.00%",      "How much higher the bull case is vs the base case"),
        ("Downside Risk  (Base Case vs Bear Case)",    downside,        "0.00%",      "How much lower the bear case is vs the base case"),
        ("Total Valuation Range  (P90 − P10)",         p90 - p10,       '$#,##0,,"B"', "The spread between best and worst outcomes — wider = more uncertainty"),
        ("Risk/Reward Ratio  (Upside ÷ Downside)",     upside/downside if downside else None, "0.00x", "Above 1.0x is favourable — more upside than downside"),
    ]
    for idx, (label, val, fmt, note) in enumerate(metrics):
        shade = XL_LIGHT_GOLD if idx % 2 == 0 else XL_WHITE
        _xlbl(ws, r, 2, label, bold=True)
        ws.cell(row=r, column=2).fill = _xfill(shade)
        c = ws.cell(row=r, column=3, value=val)
        c.font = _xf(bold=True, size=11); c.fill = _xfill(shade); c.border = _xborder()
        c.number_format = fmt; c.alignment = Alignment(horizontal="right")
        nc = ws.cell(row=r, column=4, value=note)
        nc.font = _xf(size=9, italic=True); nc.fill = _xfill(shade); nc.border = _xborder()
        nc.alignment = Alignment(wrap_text=True)
        r += 1


# ---------------------------------------------------------------------------
# Utility functions
# ---------------------------------------------------------------------------

def _fmt_currency(value) -> str:
    if value is None:
        return "N/A"
    v = float(value)
    if abs(v) >= 1e12:
        return f"${v/1e12:.2f}T"
    if abs(v) >= 1e9:
        return f"${v/1e9:.2f}B"
    if abs(v) >= 1e6:
        return f"${v/1e6:.2f}M"
    return f"${v:,.2f}"


def _fmt(value, is_pct: bool = False) -> str:
    if value is None:
        return "N/A"
    if is_pct:
        return f"{float(value) * 100:.2f}%"
    return _fmt_currency(value)


def _safe_float(val) -> float:
    if val is None:
        return 0.0
    return float(val)
